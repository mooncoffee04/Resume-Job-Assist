import PyPDF2
from docx import Document
import logging
from pathlib import Path
from typing import Optional, Tuple
import mimetypes
import tempfile
import os

logger = logging.getLogger(__name__)

class DocumentTextExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_doc,
            '.txt': self._extract_from_txt
        }
    
    def extract_text(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text from document
        Returns: (extracted_text, error_message)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return None, f"File not found: {file_path}"
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_types:
            return None, f"Unsupported file type: {extension}"
        
        try:
            # Get appropriate extraction method
            extract_method = self.supported_types[extension]
            text = extract_method(file_path)
            
            if not text or not text.strip():
                return None, "No text content found in document"
            
            # Clean up the text
            cleaned_text = self._clean_text(text)
            
            logger.info(f"✅ Successfully extracted {len(cleaned_text)} characters from {file_path.name}")
            return cleaned_text, None
            
        except Exception as e:
            error_msg = f"Failed to extract text from {file_path.name}: {str(e)}"
            logger.error(error_msg)
            return None, error_msg
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF is encrypted: {file_path.name}")
                    return ""
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        logger.debug(f"Extracted text from page {page_num + 1}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                if not text.strip():
                    logger.warning(f"No text extracted from PDF: {file_path.name}")
                
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
        
        return text
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        text = ""
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path.name}")
                
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise
        
        return text
    
    def _extract_from_doc(self, file_path: Path) -> str:
        """Extract text from DOC file (older Word format)"""
        # For .doc files, we'll need python-docx2txt or antiword
        # For now, return an error message
        raise NotImplementedError(
            "DOC file support not implemented yet. "
            "Please convert to DOCX or PDF format."
        )
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Text file extraction error: {e}")
                raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())  # Remove extra spaces
            if cleaned_line:  # Skip empty lines
                lines.append(cleaned_line)
        
        # Join with single newlines
        cleaned_text = '\n'.join(lines)
        
        # Remove very short lines that are likely formatting artifacts
        lines = cleaned_text.split('\n')
        meaningful_lines = []
        
        for line in lines:
            # Keep lines that are longer than 3 characters or contain important keywords
            if (len(line) > 3 or 
                any(keyword in line.lower() for keyword in 
                    ['email', 'phone', 'linkedin', 'github', '@', '.com', '.edu'])):
                meaningful_lines.append(line)
        
        return '\n'.join(meaningful_lines)
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about the document"""
        file_path = Path(file_path)
        
        info = {
            'filename': file_path.name,
            'extension': file_path.suffix.lower(),
            'size_bytes': file_path.stat().st_size if file_path.exists() else 0,
            'mime_type': mimetypes.guess_type(str(file_path))[0],
            'supported': file_path.suffix.lower() in self.supported_types
        }
        
        return info

# Global instance
document_extractor = DocumentTextExtractor()

def extract_text_from_file(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Convenience function for text extraction"""
    return document_extractor.extract_text(file_path)

if __name__ == "__main__":
    # Test the text extractor
    import tempfile
    
    # Create a test text file
    test_content = """
    John Doe
    Software Engineer
    
    Skills:
    - Python programming
    - Machine Learning
    - Data Analysis
    
    Experience:
    2+ years in software development
    
    Email: john.doe@example.com
    Phone: (555) 123-4567
    """
    
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write(test_content)
        temp_file = f.name
    
    try:
        # Test extraction
        extractor = DocumentTextExtractor()
        text, error = extractor.extract_text(temp_file)
        
        if error:
            print(f"❌ Error: {error}")
        else:
            print(f"✅ Extracted text ({len(text)} chars):")
            print("-" * 50)
            print(text)
            print("-" * 50)
            
        # Test document info
        info = extractor.get_document_info(temp_file)
        print(f"✅ Document info: {info}")
        
    finally:
        # Clean up
        os.unlink(temp_file)
        print("✅ Text extraction test completed!")